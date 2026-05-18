# -*- coding: utf-8 -*-
"""
Word 文档生成器 - 使用 docxtpl 生成教案 Word 文档

功能:
1. 加载教案模板
2. 填充 LLM 生成的数据
3. 导出 Word 文件
"""
import os
import re
from typing import Dict, Any, Optional
from datetime import datetime
from pathlib import Path

from create_logger import logger


class LessonPlanGenerator:
    """教案 Word 生成器"""
    
    def __init__(self, template_dir: str = None, output_dir: str = None):
        """
        初始化生成器
        
        Args:
            template_dir: 模板目录
            output_dir: 输出目录
        """
        project_root = Path(__file__).parent.parent
        self.template_dir = template_dir or str(project_root / "templates")
        self.output_dir = output_dir or str(project_root / "output")
        
        # 确保输出目录存在
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate(self, data: Dict[str, Any], template_name: str = "lesson_plan_template.docx") -> str:
        """
        生成教案 Word 文档
        
        Args:
            data: 教案数据，包含以下字段:
                - title: 教案标题
                - grade: 年级
                - unit: 单元
                - topic: 课题
                - duration: 课时
                - objectives: 教学目标（列表或字符串）
                - key_points: 教学重点
                - difficult_points: 教学难点
                - teaching_aids: 教学准备
                - procedures: 教学过程（列表）
                - homework: 作业布置
                - reflection: 教学反思
            template_name: 模板文件名
            
        Returns:
            生成的文件路径
        """
        template_path = os.path.join(self.template_dir, template_name)
        
        # 尝试使用 docxtpl，如果不可用则使用 python-docx
        try:
            from docxtpl import DocxTemplate
            
            if not os.path.exists(template_path):
                logger.warning(f"模板文件不存在: {template_path}，使用程序化生成")
                return self._generate_without_template(data)
            
            # 加载模板
            doc = DocxTemplate(template_path)
            
            # 处理数据格式
            context = self._prepare_context(data)
            
            # 渲染模板
            doc.render(context)
            
            # 生成输出文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            title = re.sub(r'[<>:"/\\|?*]', '', data.get("title", "教案")).replace(" ", "_")
            output_filename = f"{title}_{timestamp}.docx"
            output_path = os.path.join(self.output_dir, output_filename)
            
            # 保存文件
            doc.save(output_path)
            logger.info(f"教案已生成: {output_path}")
            
            return output_path
            
        except ImportError:
            logger.warning("docxtpl 未安装，使用 python-docx 程序化生成")
            return self._generate_without_template(data)
    
    def _prepare_context(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """准备模板上下文 - 适配教案生成模板"""
        context = {
            # 基本信息
            "school_name": data.get("school_name", ""),
            "topic_content": data.get("topic_content", data.get("title", "")),
            "teacher": data.get("teacher", ""),
            "lesson_type": data.get("lesson_type", "Reading"),
            "duration": data.get("duration", "1"),
            
            # 教学分析
            "situation_analysis": data.get("situation_analysis", ""),
            "teaching_objectives": data.get("teaching_objectives", ""),
            "key_points": data.get("key_points", ""),
            "difficult_points": data.get("difficult_points", ""),
            
            # 教法学法
            "teaching_methods": data.get("teaching_methods", ""),
            "learning_methods": data.get("learning_methods", ""),
            
            # 教学互动过程
            "lead_in_teacher": data.get("lead_in_teacher", ""),
            "lead_in_student": data.get("lead_in_student", ""),
            "lead_in_purpose": data.get("lead_in_purpose", ""),
            "new_lesson_teacher": data.get("new_lesson_teacher", ""),
            "new_lesson_student": data.get("new_lesson_student", ""),
            "new_lesson_purpose": data.get("new_lesson_purpose", ""),
            
            # 课堂小结、作业、板书、反思
            "summary": data.get("summary", ""),
            "summary_purpose": data.get("summary_purpose", ""),
            "homework": data.get("homework", ""),
            "homework_purpose": data.get("homework_purpose", ""),
            "board_design": data.get("board_design", ""),
            "board_purpose": data.get("board_purpose", ""),
            "reflection": data.get("reflection", ""),
        }
        
        return context
    
    def _generate_without_template(self, data: Dict[str, Any]) -> str:
        """不使用模板，程序化生成 Word 文档"""
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 标题
        title = doc.add_heading(data.get("title", "英语教案"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息表格
        table = doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        
        info = [
            ["年级", data.get("grade", ""), "单元", data.get("unit", "")],
            ["课题", data.get("topic", ""), "课时", data.get("duration", "1课时")],
            ["日期", data.get("date", datetime.now().strftime("%Y年%m月%d日")), "教师", data.get("teacher", "")],
            ["教学准备", data.get("teaching_aids", "多媒体课件"), "", ""],
        ]
        
        for i, row_data in enumerate(info):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                row.cells[j].text = str(cell_text)
        
        doc.add_paragraph()
        
        # 教学目标
        doc.add_heading("一、教学目标", level=1)
        objectives = data.get("objectives", [])
        if isinstance(objectives, list):
            for obj in objectives:
                doc.add_paragraph(obj, style='List Bullet')
        else:
            doc.add_paragraph(str(objectives))
        
        # 教学重难点
        doc.add_heading("二、教学重难点", level=1)
        doc.add_paragraph(f"重点：{data.get('key_points', '')}")
        doc.add_paragraph(f"难点：{data.get('difficult_points', '')}")
        
        # 教学过程
        doc.add_heading("三、教学过程", level=1)
        procedures = data.get("procedures", [])
        if isinstance(procedures, list):
            for i, proc in enumerate(procedures):
                if isinstance(proc, dict):
                    step = proc.get("step", f"步骤{i+1}")
                    content = proc.get("content", "")
                    doc.add_heading(step, level=2)
                    doc.add_paragraph(content)
                else:
                    doc.add_paragraph(str(proc))
        else:
            doc.add_paragraph(str(procedures))
        
        # 作业布置
        doc.add_heading("四、作业布置", level=1)
        doc.add_paragraph(data.get("homework", ""))
        
        # 教学反思
        doc.add_heading("五、教学反思", level=1)
        doc.add_paragraph(data.get("reflection", "（课后填写）"))
        
        # 保存
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        title_clean = re.sub(r'[<>:"/\\|?*]', '', data.get("title", "教案")).replace(" ", "_")
        output_filename = f"{title_clean}_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        
        doc.save(output_path)
        logger.info(f"教案已生成（程序化）: {output_path}")
        
        return output_path


def generate_lesson_plan(data: Dict[str, Any], template_name: str = "lesson_plan_template.docx") -> str:
    """生成教案的便捷函数"""
    generator = LessonPlanGenerator()
    return generator.generate(data, template_name)


if __name__ == "__main__":
    # 测试代码
    test_data = {
        "title": "Unit 1 Making Friends",
        "grade": "七年级",
        "unit": "Unit 1",
        "topic": "Making Friends",
        "duration": "1课时",
        "objectives": [
            "能够听懂并使用日常交际用语进行自我介绍",
            "能够正确使用 Nice to meet you 等问候语",
            "了解西方国家的见面礼仪"
        ],
        "key_points": "掌握自我介绍的基本表达",
        "difficult_points": "在真实情境中灵活运用所学句型",
        "teaching_aids": "多媒体课件、录音机、图片",
        "procedures": [
            {"step": "Step 1: Warm-up (5分钟)", "content": "播放英语歌曲，营造轻松的课堂氛围。教师用英语向学生问好。"},
            {"step": "Step 2: Presentation (15分钟)", "content": "教师展示图片，引入新词汇和句型。带领学生朗读和练习。"},
            {"step": "Step 3: Practice (15分钟)", "content": "学生两人一组进行对话练习，教师巡视指导。"},
            {"step": "Step 4: Production (8分钟)", "content": "请学生上台表演对话，全班点评。"},
            {"step": "Step 5: Summary (2分钟)", "content": "总结本节课所学内容，布置作业。"}
        ],
        "homework": "1. 抄写本课生词5遍\n2. 完成课后练习第1-3题\n3. 和同学用英语互相介绍",
        "reflection": ""
    }
    
    output_path = generate_lesson_plan(test_data)
    print(f"测试教案生成完成: {output_path}")
